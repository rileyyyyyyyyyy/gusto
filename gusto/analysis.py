import re
import io
import os
import logging
import tempfile
import unicodedata
from abc import ABC, abstractmethod
from dataclasses import dataclass
from datetime import datetime
from typing import Optional, Any

import magic
from PyPDF2 import PdfReader
from docx import Document
from docx.document import Document as DocxDocument

from gusto.adapter import PDFConverterAdapter

# suppress INFO logs from pdf2docx
logging.getLogger().setLevel(logging.ERROR)


class PDFOpenError(Exception):
    pass


class MetaDataReadError(Exception):
    pass


@dataclass
class DocumentAnalysis:
    word_count: int
    char_count: int
    page_count: int  # used for actual pages, or lines in plain text
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    producer: Optional[str] = None
    created: Optional[str] = None
    modified: Optional[str] = None
    mime_type: Optional[str] = None


def clean_meta(value: Any) -> Optional[str]:
    if value is None:
        return None
    try:
        if isinstance(value, bytes):
            return value.decode("utf-8", errors="ignore").strip()
        return str(value).strip()
    except Exception as e:
        raise MetaDataReadError(f'Error while reading meta data: {e}')


def clean_text_for_counting(text: str) -> str:
    text = re.sub(r'\s+', ' ', text)
    text = ''.join(c for c in text if unicodedata.category(c)[0] != 'C')
    return text.strip()


class FileAnalyser(ABC):
    def __init__(self, path: str) -> None:
        self.path = path
        self.mime_type: str = magic.from_file(path, mime=True)

    @abstractmethod
    def analyse(self) -> DocumentAnalysis:
        pass

    @abstractmethod
    def _read_metadata(self) -> dict[str, Optional[str]]:
        pass

    def _get_filesystem_dates(self) -> dict[str, Optional[str]]:
        try:
            stat = os.stat(self.path)
            created_ts = getattr(stat, 'st_birthtime', stat.st_ctime)
            modified_ts = stat.st_mtime
            return {
                "created": datetime.fromtimestamp(created_ts).strftime('%Y-%m-%d %H:%M:%S'),
                "modified": datetime.fromtimestamp(modified_ts).strftime('%Y-%m-%d %H:%M:%S'),
            }
        except Exception:   # TODO: narrow the exception clause
            return {"created": None, "modified": None}


class PDFAnalyser(FileAnalyser):
    @staticmethod
    def parse_pdf_date(raw: Optional[str]) -> Optional[str]:
        if not raw:
            return None
        if raw.startswith("D:"):
            raw = raw[2:]
        try:
            dt: datetime = datetime.strptime(raw[:14], "%Y%m%d%H%M%S")
            return dt.strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            return raw

    def __init__(self, path: str) -> None:
        super().__init__(path)
        self.reader: Optional[PdfReader] = None
        self.pages: list[Any] = []

        try:
            with open(self.path, 'rb') as file:
                data: bytes = file.read()
                self.reader = PdfReader(io.BytesIO(data))
                self.pages = self.reader.pages
        except Exception as e:
            logging.error(f"Error opening PDF: {e}")
            raise PDFOpenError(f"Error opening PDF: {e}")

    def analyse(self) -> DocumentAnalysis:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.close()
            try:
                with PDFConverterAdapter(self.path) as adapter:
                    adapter.convert(tmp.name)

                doc: DocxDocument = Document(tmp.name)
                text: str = self._extract_text(doc)
                cleaned: str = clean_text_for_counting(text)

                word_count: int = len([w for w in cleaned.split() if any(c.isalpha() for c in w)])
                char_count: int = len(cleaned.replace(" ", ""))
            finally:
                os.remove(tmp.name)

        meta: dict[str, Optional[str]] = self._read_metadata()
        fallback = self._get_filesystem_dates()

        return DocumentAnalysis(
            word_count=word_count,
            char_count=char_count,
            page_count=len(self.pages),
            title=meta.get("title"),
            author=meta.get("author"),
            subject=meta.get("subject"),
            producer=meta.get("producer"),
            created=meta.get("created") or fallback["created"],
            modified=meta.get("modified") or fallback["modified"],
            mime_type=self.mime_type,
        )

    def _extract_text(self, doc: DocxDocument) -> str:
        text_parts: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return ' '.join(text_parts)

    def _read_metadata(self) -> dict[str, Optional[str]]:
        if not self.reader:
            return {}
        metadata: dict[str, Any] = dict(self.reader.metadata or {})
        return {
            "title": clean_meta(metadata.get('/Title')),
            "author": clean_meta(metadata.get('/Author')),
            "subject": clean_meta(metadata.get('/Subject')),
            "producer": clean_meta(metadata.get('/Producer')),
            "created": self.parse_pdf_date(clean_meta(metadata.get('/CreationDate'))),
            "modified": self.parse_pdf_date(clean_meta(metadata.get('/ModDate'))),
        }


class DocumentAnalyser(FileAnalyser):
    def analyse(self) -> DocumentAnalysis:
        try:
            doc: DocxDocument = Document(self.path)
        except Exception as e:
            raise PDFOpenError(f"Error opening Word document: {e}")

        text = self._extract_text(doc)
        cleaned: str = clean_text_for_counting(text)
        word_count: int = len([w for w in cleaned.split() if any(c.isalpha() for c in w)])
        char_count: int = len(cleaned.replace(" ", ""))
        fallback = self._get_filesystem_dates()

        return DocumentAnalysis(
            word_count=word_count,
            char_count=char_count,
            page_count=len(doc.paragraphs),
            title=None,
            author=None,
            subject=None,
            producer=None,
            created=fallback["created"],
            modified=fallback["modified"],
            mime_type=self.mime_type,
        )

    def _extract_text(self, doc: DocxDocument) -> str:
        text_parts: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)
        return ' '.join(text_parts)

    def _read_metadata(self) -> dict[str, Optional[str]]:
        # TODO: get docx metadata instead of relying on fallback.
        return {}


class TextAnalyser(FileAnalyser):
    def analyse(self) -> DocumentAnalysis:
        try:
            with open(self.path, 'r', encoding='utf-8') as f:
                lines: list[str] = f.readlines()
                text: str = ''.join(lines)
        except Exception as e:
            raise PDFOpenError(f"Error reading text file: {e}")

        cleaned: str = clean_text_for_counting(text)
        word_count: int = len([w for w in cleaned.split() if any(c.isalpha() for c in w)])
        char_count: int = len(cleaned.replace(" ", ""))
        fallback = self._get_filesystem_dates()

        return DocumentAnalysis(
            word_count=word_count,
            char_count=char_count,
            page_count=len(lines),
            created=fallback["created"],
            modified=fallback["modified"],
            mime_type=self.mime_type,
        )

    def _read_metadata(self) -> dict[str, Optional[str]]:
        return {}


# TODO: add support for video, audio, and more application/* files


class AnalyserFactory:
    _registry: dict[str, type[FileAnalyser]] = {
        'application/pdf': PDFAnalyser,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DocumentAnalyser,
    }

    @staticmethod
    def get_analyser(path: str) -> FileAnalyser:
        mime_type = magic.from_file(path, mime=True)

        if mime_type in AnalyserFactory._registry:
            return AnalyserFactory._registry[mime_type](path)

        if mime_type.startswith("text/"):
            return TextAnalyser(path)

        raise ValueError(f"Unsupported MIME type: {mime_type}")
